import os
import sys
from datetime import date

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".tmp_docx_deps"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(ROOT, "reports")
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "LeonEd_Backend_Codebase_Audit_2026-09-05.docx")

NAVY = "17365D"
PALE_BLUE = "EAF2F8"
PALE_GRAY = "F5F6F7"
LIGHT_GRAY = "D9D9D9"
WHITE = "FFFFFF"
BLACK = RGBColor(0, 0, 0)
RED = RGBColor(155, 28, 28)
ORANGE = RGBColor(180, 95, 6)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LIGHT_GRAY, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(text, style=style)
    return p


def add_number(doc, text):
    return doc.add_paragraph(text, style="List Number")


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header_row = table.rows[0]
    repeat_table_header(header_row)
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, 120, 115, 120, 115)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(font_size)
        if widths:
            cell.width = Inches(widths[i])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        cant_split(table.rows[-1])
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, PALE_BLUE)
        for i, value in enumerate(values):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cell.width = Inches(widths[i])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_finding(doc, finding_id, title, severity, evidence, impact, recommendation):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    keep_with_next(p)
    run = p.add_run(f"{finding_id} {title}")
    run.font.color.rgb = BLACK

    sev = doc.add_paragraph()
    sev.paragraph_format.space_after = Pt(4)
    label = sev.add_run(f"Severity  {severity}")
    label.bold = True
    label.font.color.rgb = RED if severity == "Critical" else ORANGE if severity == "High" else BLACK

    add_body(doc, f"Evidence. {evidence}")
    add_body(doc, f"Impact. {impact}")
    add_body(doc, f"Recommendation. {recommendation}")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = BLACK
    title.paragraph_format.space_after = Pt(16)

    for style_name, size, before, after in (
        ("Heading 1", 17, 18, 8),
        ("Heading 2", 12.5, 12, 5),
        ("Heading 3", 11, 9, 4),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        styles[style_name].font.name = "Aptos"
        styles[style_name].font.size = Pt(10.5)
        styles[style_name].font.color.rgb = BLACK
        styles[style_name].paragraph_format.space_after = Pt(3)


def build_report():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    footer = section.footer
    add_page_number(footer.paragraphs[0])

    cover_spacer = doc.add_paragraph()
    cover_spacer.paragraph_format.space_after = Pt(70)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("LeonEd Backend Codebase Audit")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("Security functional reliability and promotion workflow assessment")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLACK

    metadata = [
        ("Repository", "D:\\leoned\\backend"),
        ("Reviewed revision", "f3e5e1563f434b7a6911f1b6a1980582ef23005d"),
        ("Revision date", "1 September 2026"),
        ("Assessment date", "5 September 2026"),
        ("Primary focus", "Production risks and the promotion graduation and student departure workflows"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f"{label}  ")
        r1.bold = True
        p.add_run(value)

    doc.add_paragraph().paragraph_format.space_after = Pt(42)
    p = doc.add_paragraph()
    p.add_run("Overall conclusion  ").bold = True
    p.add_run(
        "The backend compiles, and Prisma accepts the current schema, but the reviewed revision is not ready for production use. "
        "Five critical findings permit token forgery or disclosure, unauthorized real time subscriptions, payment state manipulation when configuration is incomplete, or destructive promotion outcomes. "
        "The promotion feature is incomplete and unsafe for bulk year end processing."
    )

    doc.add_page_break()

    add_heading(doc, "Executive Summary", 1)
    add_body(
        doc,
        "The assessment identified 27 findings: 5 critical, 10 high, 10 medium, and 2 low. The highest priority is to close authentication and payment fail open paths, authenticate WebSocket connections, and replace the current promotion loop with a validated transactional workflow that records immutable promotion history. Until those changes are deployed and tested, bulk promotion should remain disabled."
    )

    add_table(
        doc,
        ["Severity", "Count", "Release position"],
        [
            ["Critical", "5", "Release blocker"],
            ["High", "10", "Fix before production"],
            ["Medium", "10", "Schedule before broad rollout"],
            ["Low", "2", "Engineering hygiene"],
        ],
        widths=[1.4, 0.8, 4.7],
        font_size=9,
    )

    add_heading(doc, "Immediate Decisions", 2)
    for item in (
        "Disable or hide the bulk promotion endpoint until the transaction and mapping defects are fixed.",
        "Require JWT, Paystack, cron, and super administrator bootstrap secrets at startup. Production must fail to start when any required secret is absent.",
        "Remove password reset tokens from API responses and rotate any credentials that may have been exposed through logs or clients.",
        "Reject unauthenticated Socket.IO connections and derive room membership from the verified token rather than client supplied identifiers.",
        "Block release on automated tests for tenant isolation, role authorization, payments, results, attendance, graduation, and promotion edge cases.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Promotion Feature Verdict", 2)
    add_body(
        doc,
        "Verdict. The feature works only for a narrow single mapping in which both class identifiers belong to the school and no later mapping uses the target as a source. It does not work safely as a general bulk promotion feature. The common ascending sequence of class mappings can promote the same students multiple levels in one request. It also does not record promotion events, validate academic sessions or class levels, or guarantee all or nothing execution."
    )

    add_heading(doc, "Scope and Verification", 1)
    add_body(
        doc,
        "The review covered the Express routes, controllers, services, validation schemas, authentication and tenant middleware, Prisma schema and migrations, payment integration, report generation, attendance, results, scores, user lifecycle, and promotion services. The repository contains 107 TypeScript and Prisma source files in the reviewed scope."
    )
    add_table(
        doc,
        ["Check", "Result", "Meaning"],
        [
            ["TypeScript production build", "Passed", "npm run build completed without compiler errors"],
            ["Prisma schema validation", "Passed", "The current schema file is syntactically valid"],
            ["Dependency vulnerability audit", "Failed", "npm audit reported 7 vulnerable packages: 4 high and 3 moderate"],
            ["Automated test suite", "Unavailable", "No test script or repository test suite was found"],
            ["Promotion integration test", "Not run against live data", "The endpoint mutates student and account records; conclusions are based on deterministic code path analysis"],
        ],
        widths=[1.8, 1.25, 3.85],
        font_size=8.6,
    )
    add_body(
        doc,
        "This is a source and configuration review, not a network penetration test. A clean build does not establish correct authorization, tenant isolation, data integrity, or runtime behavior."
    )

    doc.add_page_break()
    add_heading(doc, "Promotion Workflow Assessment", 1)
    add_body(
        doc,
        "The promotion module exposes three SchoolAdmin operations: bulk class promotion, class graduation, and marking one student as Left. School scoping at the route and source or target class lookups is present. The main defects are in batch semantics, history, migration completeness, and account revocation."
    )

    add_heading(doc, "Promotion Scenario Results", 2)
    add_table(
        doc,
        ["Scenario", "Expected", "Observed from code", "Result"],
        [
            ["One valid source to target mapping", "Move active students once", "updateMany moves active students to the target", "Conditional pass"],
            ["Class 1 to 2 followed by Class 2 to 3", "Each original cohort moves one level", "The second query also selects students just moved from Class 1", "Fail"],
            ["One invalid mapping in a batch", "Reject the batch or roll back", "The service records an error and continues; earlier and later writes remain", "Fail"],
            ["Source equals target", "Reject no op mapping", "The service counts active students as promoted without changing their class", "Fail"],
            ["Target belongs to old or wrong session", "Reject invalid academic transition", "Only school ownership is checked", "Fail"],
            ["Promotion reporting", "List promoted students by session and class", "No promotion record is written; active students are indistinguishable", "Fail"],
            ["Graduate a class", "Update students and revoke access atomically", "Student statuses update before accounts; failures can leave partial state", "Fail"],
            ["Mark a student Left on a migration built database", "Persist Left and revoke access", "Migration history never adds Left to StudentStatus", "Fail"],
            ["Use an existing refresh token after leaving", "Access remains revoked", "Refresh does not check isActive and promotion does not revoke refresh tokens", "Fail"],
            ["Cross school class mapping", "Reject", "Both classes are looked up with schoolId", "Pass"],
        ],
        widths=[1.35, 1.7, 3.0, 0.85],
        font_size=7.7,
    )

    add_heading(doc, "Why Chained Mappings Fail", 2)
    add_number(doc, "The first mapping updates every active student whose current class is Class 1 and sets the class to Class 2.")
    add_number(doc, "The second mapping queries the live Student table for every active student currently in Class 2.")
    add_number(doc, "That result now contains both the original Class 2 cohort and the students moved in step 1, so both cohorts move to Class 3.")
    add_body(
        doc,
        "The outcome depends on the order of mappings supplied by the client. An administrative batch operation must not have order dependent results unless that behavior is explicit and validated."
    )

    add_heading(doc, "Required Promotion Design", 2)
    for item in (
        "Create PromotionBatch and StudentPromotion records that store source class, target class, source session, target session, student, operator, timestamp, and outcome.",
        "Validate the complete mapping graph before writing: unique source classes, source different from target, permitted level progression, target session, graduation classes, and no ambiguous cycles.",
        "Take a snapshot of student identifiers for every source class before any move. Update only those captured identifiers inside one database transaction.",
        "Use a serializable transaction or another concurrency control so two year end jobs cannot process the same cohort concurrently.",
        "Revoke refresh tokens when students graduate or leave, and make access token middleware check the current user and school activation state for sensitive requests.",
        "Add preview and dry run responses that show counts, validation failures, and the exact cohort before confirmation.",
    ):
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "Critical Findings", 1)

    add_finding(
        doc,
        "C01",
        "Known fallback JWT secret permits token forgery",
        "Critical",
        "src/middlewares/auth.middleware.ts:16 and src/utils/jwt.ts:11 use a hard coded secret when JWT_KEY is missing. Token verification does not require the configured issuer, audience, or an explicit algorithm allowlist.",
        "Anyone who knows the source can mint accepted tokens with arbitrary roles, school identifiers, and verification claims when the environment variable is absent or misconfigured.",
        "Remove the fallback, validate configuration at startup, require issuer and audience during verification, restrict algorithms, rotate the current key, and add negative token tests."
    )
    add_finding(
        doc,
        "C02",
        "Password reset secret is returned to the requester",
        "Critical",
        "src/services/auth.service.ts:306-333 generates a reset token, saves it, sends an email, and also returns resetToken and expiresAt in the public forgot password response.",
        "An unauthenticated caller who knows a registered email address can obtain the reset credential directly and take over the account.",
        "Return the same generic success response for all addresses without the token. Store only a hash of the token, use a short one time lifetime, revoke sessions after reset, and rate limit requests."
    )
    add_finding(
        doc,
        "C03",
        "WebSocket clients choose arbitrary user and school rooms",
        "Critical",
        "src/index.ts:179-195 accepts every Socket.IO connection and trusts userId and schoolId supplied in the register event before joining notification rooms.",
        "An unauthenticated client can subscribe to another user or school's real time events and receive private announcements or workflow notifications.",
        "Authenticate during the Socket.IO handshake, derive user and school identifiers from the verified token, authorize every room join, and reject client supplied identity fields."
    )
    add_finding(
        doc,
        "C04",
        "Payment and cron protections fail open when secrets are absent",
        "Critical",
        "src/controllers/payment.controller.ts:61-80 explicitly skips webhook signature verification when PAYSTACK_SECRET_KEY is not set. Lines 128-136 allow the cron route when CRON_SECRET is absent. The public webhook passes metadata schoolId and planId into subscription updates.",
        "A configuration omission makes payment state changes and account suspension or reactivation reachable without authentication. A forged charge.success payload could assign an active plan to a chosen school.",
        "Fail startup when production secrets are absent, fail closed in each handler, remove GET access for cron, use constant time signature comparison over the raw request bytes, and authenticate the scheduler."
    )
    add_finding(
        doc,
        "C05",
        "Sequential promotion mappings can move students multiple levels",
        "Critical",
        "src/services/promotion.service.ts:13-58 loops over mappings and runs updateMany against the current class membership after every prior mapping.",
        "A normal ascending class map can move a cohort through several classes in one request. This is a direct academic record integrity failure affecting every active student in the source cohort.",
        "Disable the endpoint until mappings are prevalidated, source cohorts are snapshotted before changes, and all updates plus history records execute in one transaction."
    )

    doc.add_page_break()
    add_heading(doc, "High Findings", 1)
    high_findings = [
        (
            "H01", "Deactivated students can keep using tokens",
            "src/services/auth.service.ts:243-263 refreshes any unexpired stored refresh token without checking user.isActive, school.isActive, or subscription status. src/middlewares/auth.middleware.ts trusts token claims without loading the account. Promotion only sets isActive false and leaves refresh tokens intact.",
            "Graduated or departed students can continue to obtain access tokens until the refresh token expires, contrary to the stated requirement that they can no longer log in.",
            "Revoke refresh tokens during every deactivation, check current account and school state on refresh, and introduce token versioning or session records for immediate revocation."
        ),
        (
            "H02", "Default student password is predictable and students cannot change it",
            "src/services/student.service.ts uses Student@123! when an administrator omits a password. src/services/auth.service.ts rejects password changes for the Student role.",
            "Accounts created with the default share a known permanent password until an administrator intervenes. Admission numbers are visible identifiers and are also accepted at login.",
            "Generate a unique cryptographic temporary credential or one time activation link, require a password change at first sign in, and permit secure student password changes."
        ),
        (
            "H03", "Promotion and graduation are not atomic or session aware",
            "src/services/promotion.service.ts performs class checks, student updates, and user deactivations as separate operations without prisma.$transaction. It checks school ownership but not academic session, class level, current session, duplicate mappings, cycles, or source equals target.",
            "Partial failures leave students and accounts in inconsistent states, and administrators can move students backward, into an old session, or count a no op as a promotion.",
            "Validate the full batch first and execute the snapshot updates, account changes, and audit writes in a single transaction."
        ),
        (
            "H04", "No promotion history exists for audit or reporting",
            "The Prisma schema stores only Student.classId and Student.status. The promotion service writes no event or prior class record. src/services/report.service.ts:304-338 labels a status report as including promoted students but can only return Active, Graduated, or Left.",
            "The system cannot prove who was promoted, from which class, into which session, by whom, or when. The promoted student report required in leoned req.txt:50 cannot be produced.",
            "Add immutable promotion batch and student event tables, then drive reports and rollback controls from those records."
        ),
        (
            "H05", "Migration history omits the Left student status",
            "prisma/schema.prisma:31-37 includes Left, but no SQL migration adds Left to the StudentStatus enum created in prisma/migrations/20260606014014_init/migration.sql.",
            "A database created with npm's configured migrate deploy path can reject mark-left updates and status report queries even though Prisma schema validation passes.",
            "Add and test a forward migration that alters StudentStatus to include Left. Rebuild a clean database from migrations in CI."
        ),
        (
            "H06", "Teachers can change principal level result metadata",
            "src/routes/result.routes.ts:496 permits Teacher access to PATCH /metadata/:resultId. src/controllers/result.controller.ts does not pass the caller identity or role, and src/services/result.service.ts accepts adminComment and principalsRemark without assignment or form teacher checks.",
            "Any teacher in the school who obtains a result identifier can alter principal remarks and other students' report metadata.",
            "Split teacher and administrator fields into separate endpoints or enforce field level authorization and class assignment checks in the service."
        ),
        (
            "H07", "Attendance writes do not verify that each student belongs to the class and school",
            "src/services/attendance.service.ts:119-165 validates the class and form teacher, then upserts every request studentId without checking that the student is active, belongs to schoolId, or belongs to classId.",
            "A privileged caller can create internally inconsistent attendance rows and, with a known UUID, reference a student from another tenant.",
            "Load and compare the complete expected class roster, reject unknown or duplicate identifiers, and enforce tenant consistent composite relationships where practical."
        ),
        (
            "H08", "Score writes allow inconsistent student class term and session relationships",
            "src/services/score.service.ts validates the requested class and subject in the school and separately validates the student in the school. It does not require student.classId to equal request.classId or verify that termId and academicSessionId belong to the school and to each other.",
            "Scores can be attached to the wrong class or session, producing incorrect rankings, results, reports, and historical records.",
            "Resolve the term and session server side, verify the student class membership and class subject assignment, and never accept redundant academicSessionId from the client."
        ),
        (
            "H09", "Report authorization exposes financial and personnel data too broadly",
            "src/routes/report.routes.ts:8 grants SchoolAdmin, Bursar, and Teacher access to every report route, including revenue, outstanding fees, full enrollment, student status, and staff contact details.",
            "Teachers can retrieve financial debt and revenue information and staff contact data beyond normal teaching duties.",
            "Define permissions per report. Limit revenue and fee reports to finance and administration roles, and scope teacher reports to assigned classes and subjects."
        ),
        (
            "H10", "Installed dependencies include current high severity advisories",
            "npm audit on 5 September 2026 reported 7 vulnerable packages: 4 high and 3 moderate. High findings affected brace-expansion, fast-uri, js-yaml, and socket.io-parser. Moderate findings affected express through qs, body-parser, and qs.",
            "The vulnerable packages include denial of service, host confusion or SSRF relevant parsing defects, and Socket.IO memory exhaustion.",
            "Update the lockfile with compatible patched versions, rerun the build and regression tests, and enforce npm audit thresholds in CI while reviewing actual reachability."
        ),
    ]
    for fid, title_text, evidence, impact, recommendation in high_findings:
        add_finding(doc, fid, title_text, "High", evidence, impact, recommendation)

    doc.add_page_break()
    add_heading(doc, "Medium Findings", 1)
    medium_findings = [
        (
            "M01", "Paystack signature calculation does not preserve raw request bytes",
            "src/index.ts parses JSON globally before the webhook. src/controllers/payment.controller.ts:64-69 hashes JSON.stringify(req.body) instead of the exact bytes sent by Paystack.",
            "Valid signatures can fail when whitespace or byte representation differs, while attempts to work around the issue may encourage disabling verification.",
            "Capture express.raw or a verify callback only for the webhook route and compute the HMAC over that raw buffer before JSON parsing."
        ),
        (
            "M02", "Super administrator bootstrap remains permanently reusable",
            "src/services/auth.service.ts:40-46 queries for an existing SuperAdmin, but the guard that would stop another one is commented out. The public endpoint relies on one shared SUPER_ADMIN_SECRET.",
            "Disclosure of the bootstrap secret permits creation of additional unrestricted administrators at any time.",
            "Make bootstrap single use, disable the endpoint after initialization, rotate the secret, and require an authenticated audited process for later administrator creation."
        ),
        (
            "M03", "Announcement lookup bypasses audience restrictions",
            "src/services/announcement.service.ts filters list results by role and audience, but getAnnouncementById only checks id and schoolId. The corresponding route allows any authenticated role.",
            "A student or staff member with an announcement UUID can read class, teacher, or specific user content not addressed to them.",
            "Apply the same audience predicate to item lookup and include the current user role, userId, and class assignment in the authorization decision."
        ),
        (
            "M04", "Unhandled errors expose internal details",
            "src/middlewares/error.middleware.ts returns err.message as the public problem detail for all errors, including status 500.",
            "Database constraint names, implementation details, or integration errors can be disclosed and used to refine attacks.",
            "Return a generic message for unexpected errors, log structured internal details with a correlation identifier, and map known validation and conflict errors explicitly."
        ),
        (
            "M05", "HTTP and authentication abuse controls are incomplete",
            "src/index.ts enables wildcard CORS and does not configure security headers or rate limiting. Login, OTP, password reset, registration, PDF generation, report export, and scan endpoints have no application rate controls.",
            "The service is more exposed to credential stuffing, OTP guessing, resource exhaustion, cross origin misuse, and automated data extraction.",
            "Use an explicit origin allowlist, add Helmet or equivalent headers, apply trusted proxy configuration, and enforce route specific rate and concurrency limits."
        ),
        (
            "M06", "Tenant related identifiers are not consistently validated",
            "Several services query Term by id without constraining academicSession.schoolId, including results, attendance, fees, bursar reports, and general reports. Query parameters often have no Zod schema.",
            "Known foreign identifiers can create cross tenant references or produce reports labeled with another school's session while returning current school data.",
            "Introduce reusable identifier resolvers that always scope through schoolId and validate class, term, session, subject, student, and teacher relationships together."
        ),
        (
            "M07", "Multi record account creation can leave orphaned data",
            "School registration creates School then User separately. Student creation creates User, optionally Parent, then Student. Teacher creation creates User then Teacher. These workflows do not use database transactions.",
            "A later uniqueness, foreign key, or email failure can leave a school, user, or parent record without the expected profile.",
            "Wrap each workflow in an interactive Prisma transaction and send email only after commit."
        ),
        (
            "M08", "CSV exports permit spreadsheet formula injection",
            "src/services/report.service.ts quotes CSV fields but does not neutralize values beginning with =, +, -, or @. Names and descriptions can be administrator or user supplied.",
            "Opening an exported report in spreadsheet software can evaluate attacker controlled formulas depending on client settings.",
            "Prefix dangerous leading characters with an apostrophe after normalizing whitespace, and add export regression tests."
        ),
        (
            "M09", "Revenue reporting is not an accounting ledger",
            "src/services/report.service.ts uses FeePayment.updatedAt to place the entire current amountPaid into a date range. FeePayment stores one cumulative row per student and term rather than payment transactions.",
            "Editing a payment moves and recounts the full cumulative amount, so historical revenue totals can be materially wrong.",
            "Create immutable payment transaction records with amount, date, method, reference, recorder, and reversal links. Aggregate revenue from those events."
        ),
        (
            "M10", "Validation misses important ranges and invariants",
            "Date schemas accept syntactically shaped but impossible dates and do not require startDate before endDate. Grading rules may overlap, leave gaps, omit grades, or have minScore greater than maxScore. Pagination accepts unbounded and invalid values.",
            "Invalid data produces 500 responses, ambiguous grades, misleading reports, or expensive database requests.",
            "Use refined Zod schemas for real dates, date order, grading coverage, unique grades, and bounded pagination. Return consistent 400 errors."
        ),
    ]
    for fid, title_text, evidence, impact, recommendation in medium_findings:
        add_finding(doc, fid, title_text, "Medium", evidence, impact, recommendation)

    doc.add_page_break()
    add_heading(doc, "Low Findings", 1)
    add_finding(
        doc,
        "L01",
        "No automated quality gate exists",
        "Low",
        "package.json defines build and runtime scripts but no test, lint, formatting, coverage, migration rebuild, or type aware security check. No repository test suite was found.",
        "Regressions in authorization and batch behavior can reach production even though TypeScript compiles.",
        "Add unit, integration, and API authorization tests; linting; coverage thresholds; a clean migration build; and dependency auditing to CI."
    )
    add_finding(
        doc,
        "L02",
        "Operational endpoints and API documentation are public by default",
        "Low",
        "Health, OpenAPI JSON, Swagger UI, webhook, and cron routes are registered without an environment based exposure policy. Swagger loads third party CDN assets.",
        "Public documentation and operational metadata improve endpoint discovery and add a third party runtime dependency.",
        "Restrict documentation in production, keep health responses minimal, self host static assets when required, and protect every operational endpoint."
    )

    add_heading(doc, "Remediation Plan", 1)
    add_heading(doc, "Phase One Release Blockers", 2)
    add_body(doc, "Complete before any production launch or year end processing.")
    add_table(
        doc,
        ["Work item", "Findings", "Acceptance condition"],
        [
            ["Authentication hardening", "C01 C02 H01 H02 M02", "No fallback secrets; reset tokens never leave email channel; inactive users cannot refresh; bootstrap is closed"],
            ["WebSocket authorization", "C03", "Unauthenticated connection fails; room identity comes from verified claims; cross tenant join tests pass"],
            ["Payment fail closed controls", "C04 M01", "Missing secrets stop startup; raw body signature tests pass; cron requires authenticated scheduler"],
            ["Promotion redesign", "C05 H03 H04 H05", "Snapshot based transaction and event history pass chained, invalid, replay, and rollback tests"],
            ["Role and tenant controls", "H06 H07 H08 H09 M03 M06", "A deny by default authorization matrix and cross tenant API tests pass"],
        ],
        widths=[1.45, 1.4, 4.1],
        font_size=8.2,
    )

    add_heading(doc, "Phase Two Reliability and Data Accuracy", 2)
    for item in (
        "Introduce immutable fee transaction records and rebuild revenue reporting.",
        "Make registration and profile creation transactional and enforce database level tenant invariants where possible.",
        "Add complete query validation, grading rule validation, pagination limits, and consistent domain error mapping.",
        "Sanitize CSV exports and restrict report access by role, assignment, and purpose.",
        "Upgrade vulnerable dependencies and verify behavior after lockfile changes.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Phase Three Engineering Controls", 2)
    for item in (
        "Run unit and integration tests in CI against a database built only from committed migrations.",
        "Add structured audit events for authentication, role changes, results, attendance, payments, promotions, graduation, and departures.",
        "Add security headers, origin restrictions, rate limits, request identifiers, redacted logs, and monitoring for unusual administrative batches.",
        "Document backup, restore, promotion rollback, and incident response procedures before the first production academic year rollover.",
    ):
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "Minimum Promotion Test Suite", 1)
    add_body(
        doc,
        "The corrected workflow should not be released until the following tests run against a disposable database created from migrations."
    )
    tests = [
        ["P01", "Single mapping", "Every original active student moves exactly once and one history row is written per student"],
        ["P02", "Ascending chain", "Each original cohort advances one level; no student appears in two source snapshots"],
        ["P03", "Descending chain", "The result matches P02 and does not depend on mapping order"],
        ["P04", "Duplicate source", "The request is rejected before any write"],
        ["P05", "Source equals target", "The request is rejected before any write"],
        ["P06", "Cycle", "The request is rejected before any write"],
        ["P07", "Wrong school", "The request returns forbidden or not found and writes nothing"],
        ["P08", "Wrong session or level", "The request is rejected with an actionable validation message"],
        ["P09", "Concurrent batches", "Only one batch processes a cohort; the other detects conflict"],
        ["P10", "Injected failure", "Student moves, account changes, and history all roll back"],
        ["P11", "Replay", "An idempotency key returns the original result without moving students again"],
        ["P12", "Graduation", "Statuses, accounts, refresh tokens, and history update atomically"],
        ["P13", "Mark Left", "Migration built database accepts Left and revokes every active session"],
        ["P14", "Promotion report", "Filters by source session, target session, class, date, operator, and outcome return accurate events"],
    ]
    add_table(doc, ["ID", "Case", "Required assertion"], tests, widths=[0.6, 1.55, 4.8], font_size=8.4)

    add_heading(doc, "Appendix Evidence Index", 1)
    add_body(doc, "Primary files used to confirm the findings are listed below for remediation and peer review.")
    evidence_rows = [
        ["Authentication", "src/middlewares/auth.middleware.ts; src/utils/jwt.ts; src/services/auth.service.ts"],
        ["Promotion", "src/routes/promotion.routes.ts; src/validations/promotion.validation.ts; src/services/promotion.service.ts"],
        ["Database", "prisma/schema.prisma; prisma/migrations/20260606014014_init/migration.sql"],
        ["Payments", "src/controllers/payment.controller.ts; src/services/payment.service.ts; src/utils/paystack.ts"],
        ["Real time", "src/index.ts; src/services/notification.service.ts"],
        ["Results and scores", "src/routes/result.routes.ts; src/services/result.service.ts; src/services/score.service.ts"],
        ["Attendance", "src/routes/attendance.routes.ts; src/services/attendance.service.ts"],
        ["Reports", "src/routes/report.routes.ts; src/controllers/report.controller.ts; src/services/report.service.ts"],
        ["Users", "src/services/student.service.ts; src/services/teacher.service.ts; src/services/bursar.service.ts"],
        ["Requirements", "leoned req.txt lines 37 38 and 50"],
        ["Dependencies", "package.json; package-lock.json; npm audit executed 5 September 2026"],
    ]
    add_table(doc, ["Area", "Evidence"], evidence_rows, widths=[1.45, 5.5], font_size=8.6)

    add_heading(doc, "Final Assessment", 1)
    add_body(
        doc,
        "The codebase has a usable service structure and consistent school filters in many common paths, and the reviewed revision builds successfully. Those strengths do not offset the current release blockers. Authentication can fail open, sensitive notification rooms are unauthenticated, payment controls depend on optional secrets, and the promotion algorithm can corrupt cohort placement. The safe release sequence is to close the critical access paths, redesign promotion with history and transactions, enforce relationship validation, and then establish automated regression tests before production deployment."
    )

    core_props = doc.core_properties
    core_props.title = "LeonEd Backend Codebase Audit"
    core_props.subject = "Security functional reliability and promotion workflow assessment"
    core_props.author = "LeonEd Engineering Review"
    core_props.keywords = "LeonEd, backend, security audit, promotion, Node.js, Prisma"
    core_props.comments = "Prepared from repository revision f3e5e1563f434b7a6911f1b6a1980582ef23005d"

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
